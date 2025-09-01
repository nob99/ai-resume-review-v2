"""
Text Extraction Service for UPLOAD-003

Provides comprehensive text extraction capabilities for resume files (PDF, DOCX, DOC).
Implements a factory pattern with specialized handlers for different file types.
"""

import logging
from abc import ABC, abstractmethod
from enum import Enum
from pathlib import Path
from typing import Dict, Optional, Any, List
import asyncio
from dataclasses import dataclass
import traceback

# Text extraction libraries
try:
    import PyPDF2
    import pdfplumber
    from docx import Document as DocxDocument
    from bs4 import BeautifulSoup
except ImportError as e:
    logging.error(f"Text extraction dependencies not installed: {e}")
    raise ImportError(
        "Text extraction dependencies missing. Run: pip install PyPDF2 pdfplumber python-docx beautifulsoup4 lxml"
    )

from app.core.file_validation import FileInfo


class ExtractionStatus(str, Enum):
    """Status of text extraction process."""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    TIMEOUT = "timeout"


@dataclass
class ExtractionResult:
    """Result of text extraction operation."""
    status: ExtractionStatus
    extracted_text: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    processing_time_seconds: Optional[float] = None
    
    @property
    def is_success(self) -> bool:
        """Check if extraction was successful."""
        return self.status == ExtractionStatus.COMPLETED and self.extracted_text is not None


class BaseTextExtractor(ABC):
    """Abstract base class for text extractors."""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    @abstractmethod
    def extract_text(self, file_path: Path) -> ExtractionResult:
        """Extract text from file."""
        pass
    
    @abstractmethod
    def supports_file_type(self, file_info: FileInfo) -> bool:
        """Check if this extractor supports the file type."""
        pass


class PDFTextExtractor(BaseTextExtractor):
    """PDF text extraction using multiple libraries for robustness."""
    
    def supports_file_type(self, file_info: FileInfo) -> bool:
        """Support PDF files."""
        return file_info.file_type == "application/pdf"
    
    def extract_text(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF using pdfplumber as primary, PyPDF2 as fallback."""
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Primary method: pdfplumber (better text extraction)
            text = self._extract_with_pdfplumber(file_path)
            if text and text.strip():
                return ExtractionResult(
                    status=ExtractionStatus.COMPLETED,
                    extracted_text=text,
                    metadata={
                        "extraction_method": "pdfplumber",
                        "file_type": "pdf"
                    },
                    processing_time_seconds=asyncio.get_event_loop().time() - start_time
                )
            
            # Fallback method: PyPDF2
            self.logger.info("pdfplumber returned empty text, trying PyPDF2 fallback")
            text = self._extract_with_pypdf2(file_path)
            if text and text.strip():
                return ExtractionResult(
                    status=ExtractionStatus.COMPLETED,
                    extracted_text=text,
                    metadata={
                        "extraction_method": "pypdf2",
                        "file_type": "pdf"
                    },
                    processing_time_seconds=asyncio.get_event_loop().time() - start_time
                )
            
            # No text extracted
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message="No extractable text found in PDF (may be scanned or image-based)",
                processing_time_seconds=asyncio.get_event_loop().time() - start_time
            )
            
        except Exception as e:
            self.logger.error(f"PDF extraction failed for {file_path}: {str(e)}")
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message=f"PDF extraction error: {str(e)}",
                processing_time_seconds=asyncio.get_event_loop().time() - start_time
            )
    
    def _extract_with_pdfplumber(self, file_path: Path) -> str:
        """Extract text using pdfplumber."""
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}\n")
                except Exception as e:
                    self.logger.warning(f"Failed to extract page {page_num}: {e}")
                    continue
        
        return "\n".join(text_parts)
    
    def _extract_with_pypdf2(self, file_path: Path) -> str:
        """Extract text using PyPDF2 as fallback."""
        text_parts = []
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}\n")
                except Exception as e:
                    self.logger.warning(f"PyPDF2 failed to extract page {page_num}: {e}")
                    continue
        
        return "\n".join(text_parts)


class DocxTextExtractor(BaseTextExtractor):
    """DOCX text extraction with formatting preservation."""
    
    def supports_file_type(self, file_info: FileInfo) -> bool:
        """Support DOCX files."""
        return file_info.file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    def extract_text(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX file."""
        start_time = asyncio.get_event_loop().time()
        
        try:
            document = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs with basic formatting
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text_parts)
            
            if not extracted_text.strip():
                return ExtractionResult(
                    status=ExtractionStatus.FAILED,
                    error_message="No text content found in DOCX file",
                    processing_time_seconds=asyncio.get_event_loop().time() - start_time
                )
            
            return ExtractionResult(
                status=ExtractionStatus.COMPLETED,
                extracted_text=extracted_text,
                metadata={
                    "extraction_method": "python-docx",
                    "file_type": "docx",
                    "paragraph_count": len(document.paragraphs),
                    "table_count": len(document.tables)
                },
                processing_time_seconds=asyncio.get_event_loop().time() - start_time
            )
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message=f"DOCX extraction error: {str(e)}",
                processing_time_seconds=asyncio.get_event_loop().time() - start_time
            )


class DocTextExtractor(BaseTextExtractor):
    """Legacy DOC file text extraction."""
    
    def supports_file_type(self, file_info: FileInfo) -> bool:
        """Support DOC files."""
        return file_info.file_type == "application/msword"
    
    def extract_text(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOC file (limited support)."""
        start_time = asyncio.get_event_loop().time()
        
        try:
            # For DOC files, we have limited options without external dependencies
            # This is a placeholder implementation that would need additional libraries
            # like python-docx2txt or antiword for full DOC support
            
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message="DOC file format not fully supported. Please convert to DOCX or PDF.",
                processing_time_seconds=asyncio.get_event_loop().time() - start_time
            )
            
        except Exception as e:
            self.logger.error(f"DOC extraction failed for {file_path}: {str(e)}")
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message=f"DOC extraction error: {str(e)}",
                processing_time_seconds=asyncio.get_event_loop().time() - start_time
            )


class TextExtractionService:
    """Main service for text extraction from uploaded files."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self._extractors: List[BaseTextExtractor] = [
            PDFTextExtractor(),
            DocxTextExtractor(),
            DocTextExtractor()
        ]
    
    async def extract_text_from_file(
        self, 
        file_path: Path, 
        file_info: FileInfo,
        timeout_seconds: int = 30
    ) -> ExtractionResult:
        """
        Extract text from uploaded file with timeout protection.
        
        Args:
            file_path: Path to the file to process
            file_info: File validation information
            timeout_seconds: Maximum processing time
            
        Returns:
            ExtractionResult with extracted text or error information
        """
        if not file_path.exists():
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message=f"File not found: {file_path}"
            )
        
        # Find appropriate extractor
        extractor = self._get_extractor_for_file(file_info)
        if not extractor:
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message=f"No extractor available for file type: {file_info.file_type}"
            )
        
        try:
            # Extract text with timeout protection
            result = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(
                    None,
                    extractor.extract_text,
                    file_path
                ),
                timeout=timeout_seconds
            )
            
            if result.is_success:
                self.logger.info(
                    f"Successfully extracted {len(result.extracted_text)} characters "
                    f"from {file_path.name} in {result.processing_time_seconds:.2f}s"
                )
            else:
                self.logger.warning(
                    f"Text extraction failed for {file_path.name}: {result.error_message}"
                )
            
            return result
            
        except asyncio.TimeoutError:
            self.logger.error(f"Text extraction timeout for {file_path.name}")
            return ExtractionResult(
                status=ExtractionStatus.TIMEOUT,
                error_message=f"Text extraction timed out after {timeout_seconds} seconds"
            )
        except Exception as e:
            self.logger.error(f"Unexpected error during text extraction: {str(e)}")
            self.logger.error(traceback.format_exc())
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                error_message=f"Unexpected extraction error: {str(e)}"
            )
    
    def _get_extractor_for_file(self, file_info: FileInfo) -> Optional[BaseTextExtractor]:
        """Find the appropriate extractor for the file type."""
        for extractor in self._extractors:
            if extractor.supports_file_type(file_info):
                return extractor
        return None
    
    def get_supported_file_types(self) -> List[str]:
        """Get list of supported file types for text extraction."""
        return [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]


# Global service instance
text_extraction_service = TextExtractionService()