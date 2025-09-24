"""Resume upload service with candidate-centric architecture."""

import io
import uuid
import logging
import hashlib
from typing import Optional, List, Dict, Any
from pathlib import Path

import PyPDF2
import docx
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.datetime_utils import utc_now
from .repository import ResumeRepository
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent.parent.parent))
from database.models.resume import Resume, ResumeStatus
from database.models.candidate import Candidate
from database.models.assignment import UserCandidateAssignment
from database.models.auth import User

logger = logging.getLogger(__name__)


class ResumeUploadService:
    """Service for handling resume uploads in candidate-centric architecture."""
    
    # Configuration
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    MIN_FILE_SIZE = 100  # 100 bytes
    ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}
    ALLOWED_MIME_TYPES = {
        'application/pdf',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain'
    }
    
    def __init__(self, db: Session):
        """Initialize service with database session."""
        self.db = db
        self.repository = ResumeRepository(db)
        self.settings = get_settings()
    
    async def upload_resume_for_candidate(
        self,
        candidate_id: uuid.UUID,  # NEW: Required candidate context
        file: UploadFile,
        uploaded_by_user_id: uuid.UUID,
        progress_callback: Optional[callable] = None
    ) -> Resume:
        """Upload a resume for a specific candidate."""
        
        try:
            # Step 1: Validate user can upload for this candidate
            await self._validate_candidate_access(candidate_id, uploaded_by_user_id)
            
            # Step 2: Validate file
            await self._validate_file(file)
            
            # Step 3: Read file content
            content = await file.read()
            file_size = len(content)
            
            # Step 4: Generate file hash for deduplication
            file_hash = hashlib.sha256(content).hexdigest()
            
            # Step 5: Check for duplicates
            existing_resume = await self.repository.check_duplicate_hash(candidate_id, file_hash)
            if existing_resume:
                logger.warning(f"Duplicate resume detected for candidate {candidate_id}")
                raise HTTPException(
                    status_code=400, 
                    detail=f"This resume has already been uploaded (version {existing_resume.version_number})"
                )
            
            # Step 6: Generate unique filename and get version number
            file_extension = Path(file.filename).suffix.lower()
            version_number = await self.repository.get_next_version_number(candidate_id)
            unique_filename = f"{candidate_id}_v{version_number}_{uuid.uuid4()}{file_extension}"
            
            # Step 7: Create database record
            resume = await self.repository.create_resume(
                candidate_id=candidate_id,  # NEW: Required candidate relationship
                uploaded_by_user_id=uploaded_by_user_id,
                original_filename=file.filename,
                stored_filename=unique_filename,
                file_hash=file_hash,
                file_size=file_size,
                mime_type=file.content_type or 'application/octet-stream',
                version_number=version_number
            )
            
            # Step 8: Update status to processing
            await self.repository.update_status(resume.id, ResumeStatus.PROCESSING)
            
            # Step 9: Extract text
            extracted_text = await self._extract_text(content, file_extension)
            
            # Step 10: Update with extracted content
            await self.repository.update_extracted_content(
                resume.id,
                extracted_text,
                word_count=len(extracted_text.split())
            )
            
            logger.info(f"Resume uploaded successfully for candidate {candidate_id}, version {version_number}")
            return resume
            
        except Exception as e:
            logger.error(f"Resume upload failed for candidate {candidate_id}: {str(e)}")
            
            # Mark as error if we have a DB record
            if 'resume' in locals():
                await self.repository.update_status(resume.id, ResumeStatus.ERROR)
            
            raise HTTPException(status_code=400, detail=str(e))
    
    async def get_candidate_resumes(
        self,
        candidate_id: uuid.UUID,
        user_id: uuid.UUID,
        limit: int = 10,
        offset: int = 0
    ) -> List[Resume]:
        """Get resumes for a candidate (with access control)."""
        
        # Check if user has access to this candidate
        await self._validate_candidate_access(candidate_id, user_id)
        
        return await self.repository.get_by_candidate(candidate_id, limit, offset)
    
    async def get_latest_resume_for_candidate(
        self,
        candidate_id: uuid.UUID,
        user_id: uuid.UUID
    ) -> Optional[Resume]:
        """Get the latest resume for a candidate."""
        
        # Check access
        await self._validate_candidate_access(candidate_id, user_id)
        
        return await self.repository.get_latest_version_for_candidate(candidate_id)
    
    async def get_resume_by_id(
        self,
        resume_id: uuid.UUID,
        user_id: uuid.UUID
    ) -> Optional[Resume]:
        """Get a specific resume by ID (with access control)."""
        
        resume = await self.repository.get_by_id(resume_id)
        if not resume:
            return None
        
        # Check if user has access to the candidate
        await self._validate_candidate_access(resume.candidate_id, user_id)
        
        return resume
    
    async def get_candidate_resume_stats(
        self,
        candidate_id: uuid.UUID,
        user_id: uuid.UUID
    ) -> Dict[str, Any]:
        """Get resume statistics for a candidate."""
        
        await self._validate_candidate_access(candidate_id, user_id)
        
        return await self.repository.get_candidate_resume_stats(candidate_id)
    
    # ========================================================================
    # PRIVATE METHODS
    # ========================================================================
    
    async def _validate_candidate_access(
        self,
        candidate_id: uuid.UUID,
        user_id: uuid.UUID
    ) -> None:
        """Validate that user has access to the candidate."""
        
        # Get user to check role
        user = self.db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Admin and senior recruiters can access all candidates
        if user.role in ['admin', 'senior_recruiter']:
            return
        
        # Junior recruiters can only access assigned candidates
        if user.role == 'junior_recruiter':
            assignment = (
                self.db.query(UserCandidateAssignment)
                .filter(
                    UserCandidateAssignment.user_id == user_id,
                    UserCandidateAssignment.candidate_id == candidate_id,
                    UserCandidateAssignment.is_active == True
                )
                .first()
            )
            
            if not assignment:
                raise HTTPException(
                    status_code=403, 
                    detail="Access denied: You are not assigned to this candidate"
                )
        else:
            raise HTTPException(status_code=403, detail="Invalid user role")
    
    async def _validate_file(self, file: UploadFile) -> None:
        """Validate file before processing."""
        
        # Check filename
        if not file.filename:
            raise ValueError("File must have a filename")
        
        # Check extension
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in self.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"File type not supported. Allowed types: {', '.join(self.ALLOWED_EXTENSIONS)}"
            )
        
        # Check MIME type
        if file.content_type and file.content_type not in self.ALLOWED_MIME_TYPES:
            raise ValueError(f"MIME type {file.content_type} not supported")
        
        # Check file size (will check actual size after reading)
        if file.size and file.size > self.MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum size is {self.MAX_FILE_SIZE / (1024*1024)}MB")
    
    async def _extract_text(self, content: bytes, file_extension: str) -> str:
        """Extract text from file content."""
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(content)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_docx_text(content)
            elif file_extension == '.txt':
                return content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Text extraction not supported for {file_extension}")
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from file: {str(e)}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content."""
        text_parts = []
        
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        text_parts = []
        
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text from document: {str(e)}")