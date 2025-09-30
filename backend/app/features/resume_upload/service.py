"""Resume upload service with candidate-centric upload and storage."""

import io
import uuid
import logging
import hashlib
from typing import Optional
from pathlib import Path

import PyPDF2
import docx
from fastapi import UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from .repository import ResumeUploadRepository
from database.models.resume import Resume, ResumeStatus
from .schemas import (
    UploadedFileV2,
    FileInfo,
    ProgressInfo
)

logger = logging.getLogger(__name__)


class ResumeUploadService:
    """Service for handling resume uploads with candidate association and version management."""
    
    # Configuration
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    MIN_FILE_SIZE = 100  # 100 bytes
    ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}
    ALLOWED_MIME_TYPES = {
        'application/pdf',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain'
    }
    
    def __init__(self, session: AsyncSession):
        """Initialize service with database session."""
        self.session = session
        self.repository = ResumeUploadRepository(session)
        self.settings = get_settings()
    
    async def upload_resume(
        self,
        candidate_id: uuid.UUID,
        file: UploadFile,
        user_id: uuid.UUID,
        progress_callback: Optional[callable] = None
    ) -> UploadedFileV2:
        """Upload a resume for a specific candidate with version management."""

        file_id = str(uuid.uuid4())

        try:
            content = await file.read()
            await self._validate_file(file, content)

            file_extension = Path(file.filename).suffix.lower()
            unique_filename = f"{file_id}{file_extension}"
            file_hash = hashlib.sha256(content).hexdigest()
            extracted_text = await self._extract_text(content, file_extension)

            db_upload = await self.repository.create_resume(
                candidate_id=candidate_id,
                uploaded_by_user_id=user_id,
                original_filename=file.filename,
                stored_filename=unique_filename,
                file_hash=file_hash,
                file_size=len(content),
                mime_type=file.content_type or 'application/octet-stream',
                extracted_text=extracted_text
            )

            await self.repository.update_status(db_upload.id, ResumeStatus.COMPLETED)

            return self._to_uploaded_file_v2(db_upload, extracted_text)

        except Exception as e:
            logger.error(f"File upload failed for {file.filename}: {str(e)}")

            # Mark as error if we have a DB record
            if 'db_upload' in locals():
                await self.repository.update_status(
                    db_upload.id,
                    ResumeStatus.ERROR,
                    str(e)
                )

            raise HTTPException(status_code=400, detail=str(e))
    
    async def _validate_file(self, file: UploadFile, content: bytes) -> None:
        """
        Validate file metadata and content.
        Consolidated validation to avoid duplicate checks.
        """
        # Check filename exists
        if not file.filename:
            raise ValueError("File must have a filename")

        # Check file extension
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in self.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"File type not supported. Allowed types: {', '.join(self.ALLOWED_EXTENSIONS)}"
            )

        # Check MIME type
        if file.content_type and file.content_type not in self.ALLOWED_MIME_TYPES:
            raise ValueError(f"MIME type {file.content_type} not supported")

        # Check file size (single check using actual content)
        content_size = len(content)
        if content_size > self.MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum size is {self.MAX_FILE_SIZE / (1024*1024)}MB")
        if content_size < self.MIN_FILE_SIZE:
            raise ValueError(f"File too small. Minimum size is {self.MIN_FILE_SIZE} bytes")

        # Check for suspicious patterns (basic content security)
        # TODO: Implement proper virus scanning
        suspicious_patterns = [b'<script', b'javascript:', b'onclick=']
        content_lower = content.lower()
        for pattern in suspicious_patterns:
            if pattern in content_lower:
                logger.warning(f"Suspicious pattern detected: {pattern}")
                # In production, might want to reject the file
    
    async def _extract_text(self, content: bytes, file_extension: str) -> str:
        """Extract text from file content based on file type."""
        try:
            if file_extension == '.pdf':
                # Extract from PDF
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                return '\n'.join(page.extract_text() for page in pdf_reader.pages)

            elif file_extension in ['.doc', '.docx']:
                # Extract from Word document
                doc = docx.Document(io.BytesIO(content))
                text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

                # Include text from tables
                for table in doc.tables:
                    text_parts.extend(
                        cell.text for row in table.rows
                        for cell in row.cells if cell.text.strip()
                    )

                return '\n'.join(text_parts)

            elif file_extension == '.txt':
                return content.decode('utf-8', errors='ignore')

            else:
                raise ValueError(f"Text extraction not supported for {file_extension}")

        except Exception as e:
            logger.error(f"Text extraction failed for {file_extension}: {str(e)}")
            raise ValueError(f"Failed to extract text from file: {str(e)}")

    def _to_uploaded_file_v2(self, db_upload: Resume, extracted_text: str) -> UploadedFileV2:
        """Convert database model to frontend-compatible schema."""

        # Map MIME type to simple file type
        file_type = 'txt'
        if db_upload.mime_type:
            if 'pdf' in db_upload.mime_type:
                file_type = 'pdf'
            elif 'msword' in db_upload.mime_type or 'document' in db_upload.mime_type:
                file_type = 'docx'

        return UploadedFileV2(
            id=str(db_upload.id),
            candidate_id=str(db_upload.candidate_id),
            filename=db_upload.original_filename,
            file=FileInfo(
                name=db_upload.original_filename,
                size=db_upload.file_size,
                type=file_type,
                lastModified=int(db_upload.uploaded_at.timestamp() * 1000)
            ),
            status=db_upload.status,
            progress=db_upload.progress or (100 if db_upload.status == ResumeStatus.COMPLETED.value else 0),
            progressInfo=ProgressInfo(
                fileId=str(db_upload.id),
                fileName=db_upload.original_filename,
                stage=db_upload.status,
                percentage=db_upload.progress or (100 if db_upload.status == ResumeStatus.COMPLETED.value else 0),
                bytesUploaded=db_upload.file_size if db_upload.status == ResumeStatus.COMPLETED.value else 0,
                totalBytes=db_upload.file_size,
                timeElapsed=0,
                estimatedTimeRemaining=0,
                speed=0,
                retryCount=0,
                maxRetries=3
            ),
            extracted_text=extracted_text or db_upload.extracted_text,
            error=None,
            startTime=int(db_upload.uploaded_at.timestamp() * 1000),
            endTime=int(db_upload.processed_at.timestamp() * 1000) if db_upload.processed_at else None
        )

    async def cancel_upload(self, file_id: uuid.UUID, user_id: uuid.UUID) -> None:
        """
        Cancel an upload if business rules allow it.

        Business Rules:
        - Upload must exist
        - User must own the upload (authorization)
        - Upload must be in cancellable state (not COMPLETED, ERROR, or CANCELLED)

        Args:
            file_id: Upload ID to cancel
            user_id: User requesting cancellation (for ownership check)

        Raises:
            ValueError: If upload not found, not owned by user, or in non-cancellable state
        """
        upload = await self.repository.get_by_id(file_id)

        # Business rule: Upload must exist and user must own it
        if not upload or upload.uploaded_by_user_id != user_id:
            raise ValueError("Upload not found")  # Security: don't reveal existence

        # Business rule: Check if cancellation is allowed by state
        non_cancellable_states = [
            ResumeStatus.COMPLETED.value,
            ResumeStatus.ERROR.value,
            ResumeStatus.CANCELLED.value
        ]

        if upload.status in non_cancellable_states:
            raise ValueError(
                f"Cannot cancel upload in '{upload.status}' state. "
                f"Only pending or uploading uploads can be cancelled."
            )

        # Perform cancellation
        await self.repository.mark_cancelled(file_id)