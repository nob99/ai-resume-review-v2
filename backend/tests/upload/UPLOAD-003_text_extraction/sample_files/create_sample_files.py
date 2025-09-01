"""
Create sample test files for UPLOAD-003 testing

Generates sample PDF, DOCX, and DOC files with realistic resume content
for comprehensive testing of text extraction functionality.
"""

import os
from pathlib import Path
from io import BytesIO
import tempfile

# Only try to import libraries if available
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def create_sample_pdf_simple(output_path: Path) -> bool:
    """Create a simple PDF file using reportlab."""
    if not REPORTLAB_AVAILABLE:
        return False
    
    try:
        # Simple PDF creation
        c = canvas.Canvas(str(output_path), pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, height - 100, "JOHN DOE")
        
        # Contact info
        c.setFont("Helvetica", 12)
        c.drawString(100, height - 130, "Software Engineer")
        c.drawString(100, height - 150, "john.doe@email.com")
        c.drawString(100, height - 170, "(555) 123-4567")
        c.drawString(100, height - 190, "LinkedIn: linkedin.com/in/johndoe")
        
        # Professional Summary
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, height - 230, "PROFESSIONAL SUMMARY")
        c.setFont("Helvetica", 12)
        summary_text = "Experienced software engineer with 5+ years developing scalable web applications."
        c.drawString(100, height - 260, summary_text)
        
        # Experience
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, height - 300, "PROFESSIONAL EXPERIENCE")
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, height - 330, "Senior Software Engineer - TechCorp Inc.")
        c.setFont("Helvetica", 10)
        c.drawString(100, height - 350, "2020 - Present")
        c.setFont("Helvetica", 12)
        c.drawString(100, height - 370, "• Led development of microservices architecture")
        c.drawString(100, height - 390, "• Implemented CI/CD pipelines reducing deployment time by 60%")
        c.drawString(100, height - 410, "• Mentored junior developers and conducted code reviews")
        
        # Education
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, height - 450, "EDUCATION")
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, height - 480, "Master of Science in Computer Science")
        c.setFont("Helvetica", 12)
        c.drawString(100, height - 500, "University of Technology - 2018")
        
        # Skills
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, height - 540, "TECHNICAL SKILLS")
        c.setFont("Helvetica", 12)
        c.drawString(100, height - 570, "Programming Languages: Python, JavaScript, TypeScript, Java")
        c.drawString(100, height - 590, "Frameworks: React, Django, Flask, Node.js")
        c.drawString(100, height - 610, "Databases: PostgreSQL, MongoDB, Redis")
        c.drawString(100, height - 630, "Cloud & DevOps: AWS, Docker, Kubernetes, Jenkins")
        
        c.save()
        return True
        
    except Exception as e:
        print(f"Error creating simple PDF: {e}")
        return False


def create_sample_pdf_complex(output_path: Path) -> bool:
    """Create a more complex PDF with multiple pages and formatting."""
    if not REPORTLAB_AVAILABLE:
        return False
    
    try:
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title = Paragraph("JANE SMITH", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Contact
        contact = Paragraph("Senior Data Scientist<br/>jane.smith@email.com<br/>(555) 987-6543<br/>GitHub: github.com/janesmith", styles['Normal'])
        story.append(contact)
        story.append(Spacer(1, 20))
        
        # Summary
        summary_title = Paragraph("PROFESSIONAL SUMMARY", styles['Heading2'])
        story.append(summary_title)
        summary_text = """
        Data scientist with 7+ years of experience in machine learning, statistical analysis, and big data processing. 
        Proven track record of building predictive models that drive business decisions and improve operational efficiency. 
        Expert in Python, R, SQL, and cloud-based analytics platforms.
        """
        summary = Paragraph(summary_text, styles['Normal'])
        story.append(summary)
        story.append(Spacer(1, 20))
        
        # Experience
        exp_title = Paragraph("PROFESSIONAL EXPERIENCE", styles['Heading2'])
        story.append(exp_title)
        
        job1_title = Paragraph("Senior Data Scientist - DataCorp Technologies (2019 - Present)", styles['Heading3'])
        story.append(job1_title)
        
        job1_details = """
        • Led cross-functional team of 8 data scientists and engineers<br/>
        • Developed machine learning models increasing customer retention by 25%<br/>
        • Built real-time recommendation system processing 1M+ events daily<br/>
        • Implemented A/B testing framework for product optimization<br/>
        • Technologies: Python, TensorFlow, Scikit-learn, Apache Spark, AWS
        """
        job1_desc = Paragraph(job1_details, styles['Normal'])
        story.append(job1_desc)
        story.append(Spacer(1, 15))
        
        job2_title = Paragraph("Data Scientist - Analytics Solutions Inc. (2017 - 2019)", styles['Heading3'])
        story.append(job2_title)
        
        job2_details = """
        • Designed predictive models for fraud detection with 95% accuracy<br/>
        • Created automated reporting dashboard using Tableau and Power BI<br/>
        • Performed statistical analysis on customer behavior patterns<br/>
        • Collaborated with product teams on feature development
        """
        job2_desc = Paragraph(job2_details, styles['Normal'])
        story.append(job2_desc)
        story.append(Spacer(1, 20))
        
        # Education
        edu_title = Paragraph("EDUCATION", styles['Heading2'])
        story.append(edu_title)
        
        edu_details = """
        <b>Ph.D. in Statistics</b> - Stanford University (2017)<br/>
        Dissertation: "Bayesian Methods for Large-Scale Time Series Analysis"<br/><br/>
        <b>Master of Science in Mathematics</b> - UC Berkeley (2014)<br/>
        Specialization: Applied Statistics and Probability
        """
        education = Paragraph(edu_details, styles['Normal'])
        story.append(education)
        story.append(Spacer(1, 20))
        
        # Skills
        skills_title = Paragraph("TECHNICAL SKILLS", styles['Heading2'])
        story.append(skills_title)
        
        skills_details = """
        <b>Programming Languages:</b> Python, R, SQL, Scala, Java<br/>
        <b>Machine Learning:</b> TensorFlow, PyTorch, Scikit-learn, XGBoost<br/>
        <b>Big Data Tools:</b> Apache Spark, Hadoop, Kafka, Elasticsearch<br/>
        <b>Cloud Platforms:</b> AWS (SageMaker, EMR, Redshift), Google Cloud Platform<br/>
        <b>Visualization:</b> Tableau, Power BI, Matplotlib, Plotly, D3.js<br/>
        <b>Databases:</b> PostgreSQL, MySQL, MongoDB, Cassandra, Redis
        """
        skills = Paragraph(skills_details, styles['Normal'])
        story.append(skills)
        story.append(Spacer(1, 20))
        
        # Certifications
        cert_title = Paragraph("CERTIFICATIONS", styles['Heading2'])
        story.append(cert_title)
        
        cert_details = """
        • AWS Certified Machine Learning - Specialty (2022)<br/>
        • Google Professional Data Engineer (2021)<br/>
        • Certified Analytics Professional (CAP) (2020)
        """
        certifications = Paragraph(cert_details, styles['Normal'])
        story.append(certifications)
        
        doc.build(story)
        return True
        
    except Exception as e:
        print(f"Error creating complex PDF: {e}")
        return False


def create_sample_docx(output_path: Path) -> bool:
    """Create a sample DOCX file with resume content."""
    if not DOCX_AVAILABLE:
        return False
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('MICHAEL JOHNSON', 0)
        title.alignment = 1  # Center alignment
        
        # Contact Info
        contact = doc.add_paragraph('Full Stack Developer\nmichael.johnson@email.com\n(555) 456-7890\nPortfolio: michaeljohnson.dev')
        contact.alignment = 1
        
        # Professional Summary
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        doc.add_paragraph(
            'Full-stack developer with 6+ years of experience building responsive web applications '
            'and RESTful APIs. Proficient in modern JavaScript frameworks, Python web development, '
            'and cloud deployment. Strong background in agile development methodologies and '
            'test-driven development.'
        )
        
        # Technical Skills
        doc.add_heading('TECHNICAL SKILLS', level=1)
        skills_table = doc.add_table(rows=4, cols=2)
        skills_table.style = 'Table Grid'
        
        skills_data = [
            ('Frontend', 'React, Vue.js, Angular, HTML5, CSS3, TypeScript'),
            ('Backend', 'Node.js, Python, Django, Flask, Express.js'),
            ('Databases', 'PostgreSQL, MongoDB, MySQL, Redis'),
            ('DevOps', 'Docker, Kubernetes, AWS, CI/CD, Jenkins')
        ]
        
        for i, (category, technologies) in enumerate(skills_data):
            skills_table.cell(i, 0).text = category
            skills_table.cell(i, 1).text = technologies
        
        # Professional Experience
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        
        # Job 1
        doc.add_heading('Senior Full Stack Developer - WebTech Solutions', level=2)
        doc.add_paragraph('March 2020 - Present')
        
        job1_responsibilities = [
            'Architected and developed microservices-based e-commerce platform handling 50K+ daily users',
            'Led frontend team of 4 developers in React-based single-page application development',
            'Implemented automated testing suite achieving 90% code coverage',
            'Optimized database queries reducing average response time by 40%',
            'Mentored junior developers and conducted technical interviews'
        ]
        
        for responsibility in job1_responsibilities:
            doc.add_paragraph(responsibility, style='List Bullet')
        
        # Job 2
        doc.add_heading('Full Stack Developer - StartupHub Inc.', level=2)
        doc.add_paragraph('June 2018 - February 2020')
        
        job2_responsibilities = [
            'Built RESTful APIs using Node.js and Express.js for mobile and web clients',
            'Developed responsive web interfaces using React and Redux',
            'Integrated third-party payment systems (Stripe, PayPal) and APIs',
            'Implemented real-time features using WebSocket connections',
            'Collaborated with UX/UI designers on user interface improvements'
        ]
        
        for responsibility in job2_responsibilities:
            doc.add_paragraph(responsibility, style='List Bullet')
        
        # Education
        doc.add_heading('EDUCATION', level=1)
        doc.add_paragraph('Bachelor of Science in Computer Science\nTech University - 2018\nGPA: 3.7/4.0')
        
        # Projects
        doc.add_heading('KEY PROJECTS', level=1)
        
        project1 = doc.add_paragraph()
        project1.add_run('Task Management SaaS Platform').bold = True
        project1.add_run(' - Full-stack web application with React frontend and Node.js backend. '
                        'Features include real-time collaboration, file uploads, and team management. '
                        'Deployed on AWS with Docker containers.')
        
        project2 = doc.add_paragraph()
        project2.add_run('Mobile-First E-commerce API').bold = True
        project2.add_run(' - RESTful API built with Python Flask supporting product catalog, '
                        'user authentication, payment processing, and order management. '
                        'Includes comprehensive API documentation and automated testing.')
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating DOCX file: {e}")
        return False


def create_text_based_files(sample_dir: Path) -> None:
    """Create simple text-based files for testing when libraries aren't available."""
    
    # Simple text resume
    simple_resume = """SARAH WILLIAMS
Marketing Manager
sarah.williams@email.com
(555) 321-9876

PROFESSIONAL SUMMARY
Results-driven marketing professional with 5+ years of experience in digital marketing, 
brand management, and campaign optimization. Proven track record of increasing brand 
awareness and driving customer acquisition through data-driven marketing strategies.

PROFESSIONAL EXPERIENCE

Digital Marketing Manager - Growth Marketing Co.
January 2020 - Present
• Developed and executed digital marketing campaigns resulting in 150% increase in lead generation
• Managed social media presence across multiple platforms with 200K+ combined followers
• Implemented marketing automation workflows improving conversion rates by 35%
• Collaborated with sales team to optimize lead qualification process

Marketing Specialist - Brand Solutions Inc.
June 2018 - December 2019
• Created content marketing strategy increasing organic website traffic by 80%
• Managed Google Ads campaigns with monthly budget of $50K+
• Conducted market research and competitive analysis for product launches
• Designed email marketing campaigns achieving 25% open rates

EDUCATION
Master of Business Administration (MBA) - Marketing Focus
Business School University - 2018

Bachelor of Arts in Communications
State College - 2016

SKILLS
• Digital Marketing: Google Ads, Facebook Ads, LinkedIn Ads
• Analytics: Google Analytics, Adobe Analytics, Tableau
• Content Management: WordPress, HubSpot, Marketo
• Design: Adobe Creative Suite, Canva, Figma
• Social Media: Hootsuite, Buffer, Sprout Social

CERTIFICATIONS
• Google Ads Certified (2022)
• HubSpot Content Marketing Certified (2021)
• Facebook Blueprint Certified (2020)
"""
    
    # Save as text file for manual conversion to other formats
    text_file_path = sample_dir / "sample_resume_text.txt"
    with open(text_file_path, 'w', encoding='utf-8') as f:
        f.write(simple_resume)
    
    print(f"Created text resume: {text_file_path}")
    
    # Instructions file
    instructions = """UPLOAD-003 Sample Files Instructions

This directory contains sample resume files for testing text extraction functionality.

Files included:
- sample_resume_text.txt: Text version of sample resume
- sample_resume_simple.pdf: Simple PDF resume (if reportlab available)
- sample_resume_complex.pdf: Complex multi-page PDF (if reportlab available)  
- sample_resume.docx: DOCX resume with tables (if python-docx available)

To create additional test files:
1. Install dependencies: pip install reportlab python-docx
2. Run this script again
3. Or manually create PDF/DOCX files using the text content as reference

Test Scenarios:
- Basic text extraction from different file formats
- Section detection and classification
- Contact information extraction
- Multi-page document handling
- Table and formatting preservation
- Error handling with corrupted files

For corrupted file testing:
- Create files with invalid headers or truncated content
- Test with password-protected documents
- Test with scanned PDFs (image-based content)
"""
    
    instructions_file = sample_dir / "README.txt"
    with open(instructions_file, 'w', encoding='utf-8') as f:
        f.write(instructions)
    
    print(f"Created instructions: {instructions_file}")


def main():
    """Create all sample files for testing."""
    # Get the directory of this script
    current_dir = Path(__file__).parent
    sample_dir = current_dir
    
    print(f"Creating sample files in: {sample_dir}")
    
    # Create directory if it doesn't exist
    sample_dir.mkdir(parents=True, exist_ok=True)
    
    created_files = []
    
    # Create simple PDF
    simple_pdf_path = sample_dir / "sample_resume_simple.pdf"
    if create_sample_pdf_simple(simple_pdf_path):
        created_files.append("sample_resume_simple.pdf")
        print(f"✓ Created simple PDF: {simple_pdf_path}")
    else:
        print("✗ Could not create simple PDF (reportlab not available)")
    
    # Create complex PDF
    complex_pdf_path = sample_dir / "sample_resume_complex.pdf"
    if create_sample_pdf_complex(complex_pdf_path):
        created_files.append("sample_resume_complex.pdf")
        print(f"✓ Created complex PDF: {complex_pdf_path}")
    else:
        print("✗ Could not create complex PDF (reportlab not available)")
    
    # Create DOCX file
    docx_path = sample_dir / "sample_resume.docx"
    if create_sample_docx(docx_path):
        created_files.append("sample_resume.docx")
        print(f"✓ Created DOCX: {docx_path}")
    else:
        print("✗ Could not create DOCX file (python-docx not available)")
    
    # Always create text-based files and instructions
    create_text_based_files(sample_dir)
    created_files.extend(["sample_resume_text.txt", "README.txt"])
    
    print(f"\n✓ Sample file creation completed!")
    print(f"Created {len(created_files)} files:")
    for filename in created_files:
        print(f"  - {filename}")
    
    if not REPORTLAB_AVAILABLE:
        print("\n📝 Note: Install 'reportlab' to generate PDF files")
    if not DOCX_AVAILABLE:
        print("📝 Note: Install 'python-docx' to generate DOCX files")
    
    print("\n🧪 Files are ready for UPLOAD-003 testing!")


if __name__ == "__main__":
    main()